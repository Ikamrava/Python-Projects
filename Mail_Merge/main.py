import docx

PLACE_HOLDER = "[name]"


with open("Mail_Merge/Input/Names/invited_names.txt") as names:
   name_list = names.readlines()


template_doc = docx.Document("Mail_Merge/Input/Letters/starting_letter.docx")


for name in name_list:
   stripped_name = name.strip()


   new_doc = docx.Document()


   for paragraph in template_doc.paragraphs:
       new_paragraph = paragraph.text.replace(PLACE_HOLDER, stripped_name)
       new_doc.add_paragraph(new_paragraph)


   new_doc.save(f"Mail_Merge/Output/ReadyToSend/letter_for_{stripped_name}.docx")
